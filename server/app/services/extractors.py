"""
app/services/extractors.py
Text extraction for each supported file type.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from loguru import logger


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct extractor based on filename extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    elif ext in (".txt", ".csv"):
        return _extract_text_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    """pdfplumber primary; pytesseract fallback for scanned PDFs."""
    import pdfplumber

    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    if not text_parts:
        logger.info("pdfplumber returned no text — attempting OCR fallback.")
        text_parts = _ocr_pdf(data)

    return "\n".join(text_parts)


def _ocr_pdf(data: bytes) -> list[str]:
    """pytesseract OCR fallback for scanned PDFs."""
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image

        doc = fitz.open(stream=data, filetype="pdf")
        results = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            results.append(pytesseract.image_to_string(img))
        return results
    except Exception as e:
        logger.error(f"OCR fallback failed: {e}")
        return []


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


# ── TXT / CSV ─────────────────────────────────────────────────────────────────

def _extract_text_csv(data: bytes) -> str:
    import chardet

    detected = chardet.detect(data)
    encoding = detected.get("encoding") or "utf-8"
    return data.decode(encoding, errors="replace")
