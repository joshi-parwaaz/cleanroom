"""
app/services/reconstructor.py
Rebuilds the original file format with entities replaced by pseudonyms.

KEY FIX: str.replace() in a loop was the root cause of all output corruption.
  - With 896 entities it ran thousands of global scans on already-mutated text
  - Short entity texts ("me", "al") matched inside words like "medical", "local"
  - Each iteration re-scanned previously placed pseudonyms, cascading into artifacts
    like "SaKenneth Gibsonah Holt" and "[Kenneth GibsonEDACTED_PYNB]"

SOLUTION: Single-pass re.sub with a compiled alternation of all entity strings.
  - Longest strings listed first -> greedy match prevents partial-word hits
  - re.sub visits each character ONCE -> replaced text is NEVER re-scanned
  - No cascading, no substring corruption, deterministic output
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import List

from app.services.reconciler import ReconciledEntity


# ─── Shared helpers ───────────────────────────────────────────────────────────

def _build_replacement_map(entities: List[ReconciledEntity], registry) -> dict:
    """
    Build a deduplicated {original_text: pseudonym} map.
    Longer originals come first so the regex alternation matches greedily
    and 'Sarah Holt' wins over 'Sarah' at the same position.
    """
    mapping: dict[str, str] = {}
    for ent in sorted(entities, key=lambda e: len(e.text), reverse=True):
        if ent.text.strip() and ent.text not in mapping:
            mapping[ent.text] = registry.get_or_create(ent.text, ent.entity_type)
    return mapping


def _single_pass_replace(text: str, replacement_map: dict) -> str:
    """
    Replace every entity string in ONE re.sub pass — zero cascading, zero corruption.

    Two layers of protection against partial-word matches:
      1. The reconciler already dropped mid-word NER spans.
      2. Here we wrap each key with word-boundary guards ((?<!\\w) / (?!\\w))
         when the key starts/ends with a word character.  This means even
         if a short entity text slips through, it can ONLY match at a
         genuine word boundary, never inside "vascular", "medical", etc.

    re.sub visits each character exactly once — replaced text is never
    rescanned, so pseudonym A cannot cascade into pseudonym B.
    """
    if not replacement_map:
        return text
    pairs = [(k, v) for k, v in replacement_map.items() if k.strip()]
    if not pairs:
        return text

    parts = []
    for key, _ in pairs:
        escaped = re.escape(key)
        # Add zero-width word-boundary assertions only when needed
        prefix = r"(?<!\w)" if key[0].isalnum() or key[0] == "_" else ""
        suffix = r"(?!\w)"  if key[-1].isalnum() or key[-1] == "_" else ""
        parts.append(prefix + escaped + suffix)

    pattern = re.compile("|".join(parts))
    lut = dict(pairs)
    # match.group(0) is always the original key since we only match exact strings
    return pattern.sub(lambda m: lut[m.group(0)], text)



# ─── Offset-based replacement (demo + TXT/CSV) ────────────────────────────────

def apply_replacements(text: str, entities: List[ReconciledEntity], registry) -> str:
    """
    Right-to-left offset-based replacement.
    Used by: demo endpoint, TXT/CSV reconstruction.
    Requires that `text` is the SAME string used during entity detection
    so character offsets are valid.
    """
    result = list(text)
    for ent in sorted(entities, key=lambda e: e.start, reverse=True):
        pseudo = registry.get_or_create(ent.text, ent.entity_type)
        result[ent.start:ent.end] = list(pseudo)
    return "".join(result)


# ─── Main dispatch ────────────────────────────────────────────────────────────

def reconstruct(
    original_bytes: bytes,
    filename: str,
    entities: List[ReconciledEntity],
    registry,
) -> bytes:
    """Reconstruct file in its original format with entities replaced."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _reconstruct_pdf(original_bytes, entities, registry)
    elif ext == ".docx":
        return _reconstruct_docx(original_bytes, entities, registry)
    else:
        # TXT / CSV — offset-based replacement preserving exact encoding
        import chardet
        detected = chardet.detect(original_bytes)
        enc = detected.get("encoding") or "utf-8"
        text = original_bytes.decode(enc, errors="replace")
        cleaned = apply_replacements(text, entities, registry)
        return cleaned.encode(enc)


# ─── PDF reconstruction ───────────────────────────────────────────────────────

def _reconstruct_pdf(data: bytes, entities: List[ReconciledEntity], registry) -> bytes:
    """
    Extract text page-by-page with fitz, apply single-pass replacement,
    write a clean new PDF.

    Uses ONLY Python-layer fitz ops:
      fitz.open / page.get_text / new_page / insert_textbox / save
    No search_for, no apply_redactions — those C-level calls hang on
    Windows PyMuPDF 1.24.x regardless of Python try/except wrappers.
    """
    import fitz

    replacement_map = _build_replacement_map(entities, registry)

    original_doc = fitz.open(stream=data, filetype="pdf")
    new_doc = fitz.open()

    for orig_page in original_doc:
        page_text  = orig_page.get_text()
        clean_text = _single_pass_replace(page_text, replacement_map)

        w = orig_page.rect.width  or 595
        h = orig_page.rect.height or 842
        new_page = new_doc.new_page(width=w, height=h)

        # Use insert_textbox with a margin rect so text wraps and never clips.
        # rect shrunk by 50pt on all sides gives a clean readable margin.
        text_rect = fitz.Rect(50, 50, w - 50, h - 50)
        new_page.insert_textbox(
            text_rect,
            clean_text,
            fontsize=9,
            fontname="cour",
            lineheight=1.4,   # breathing room between lines
            expandtabs=4,
        )

    original_doc.close()

    buf = io.BytesIO()
    new_doc.save(buf)
    return buf.getvalue()


# ─── DOCX reconstruction ──────────────────────────────────────────────────────

def _reconstruct_docx(data: bytes, entities: List[ReconciledEntity], registry) -> bytes:
    from docx import Document

    replacement_map = _build_replacement_map(entities, registry)

    doc = Document(io.BytesIO(data))

    def clean(text: str) -> str:
        return _single_pass_replace(text, replacement_map)

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = clean(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = clean(run.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── NOT CALLED — kept for reference only ────────────────────────────────────

def _reconstruct_pdf_visual_UNSAFE(data: bytes, entities: List[ReconciledEntity], registry) -> bytes:
    """
    Visual redaction via draw_rect + insert_text.
    page.search_for() is a C-level MuPDF call that hangs on Windows
    PyMuPDF 1.24.x. Do not call from production code.
    """
    import fitz
    doc = fitz.open(stream=data, filetype="pdf")
    rm  = _build_replacement_map(entities, registry)
    for page in doc:
        for orig, pseudo in rm.items():
            try:
                for rect in page.search_for(orig, quads=False):
                    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                    page.insert_text((rect.x0, rect.y1 - 1), pseudo, fontsize=9)
            except Exception:
                pass
    buf = io.BytesIO()
    doc.save(buf, garbage=3, deflate=True)
    return buf.getvalue()